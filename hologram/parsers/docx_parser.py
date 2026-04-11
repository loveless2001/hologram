"""DOCX parser for Hologram ingestion.

Extracts text paragraphs, embedded images, and basic structure from .docx
files using python-docx. Produces the same ParsedDocument contract as the
PDF parser so the downstream ingestion path is identical.
"""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import List, Optional

from .types import ParsedDocument, ParsedImage, ParsedPage


def _doc_id_for_path(path: Path) -> str:
    digest = hashlib.blake2b(str(path.resolve()).encode("utf-8"), digest_size=8).hexdigest()
    return f"doc:{digest}"


def parse_docx(
    path: str,
    image_output_dir: Optional[str] = None,
    extract_images: bool = True,
) -> ParsedDocument:
    """Parse a DOCX file into pages of text and extracted images.

    Requires `python-docx` at runtime. The dependency is optional so the
    wider project can still run without DOCX ingestion support installed.

    DOCX files don't have physical pages, so we treat the entire document
    as a single logical page. Section breaks could be used for finer splits
    in a future version.
    """
    try:
        from docx import Document  # type: ignore
        from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "DOCX ingestion requires python-docx. Install with `pip install python-docx`."
        ) from exc

    docx_path = Path(path).expanduser().resolve()
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    doc = Document(str(docx_path))
    doc_id = _doc_id_for_path(docx_path)

    output_dir = Path(image_output_dir).expanduser().resolve() if image_output_dir else (
        docx_path.parent / ".hologram_extracted_images" / doc_id
    )

    # Extract text from paragraphs, preserving block structure
    blocks: List[dict] = []
    text_parts: List[str] = []
    for block_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        text_parts.append(text)
        blocks.append({
            "block_index": block_idx,
            "bbox": None,  # DOCX doesn't have spatial layout
            "text": text,
            "style": para.style.name if para.style else None,
        })

    # Extract embedded images from document relationships
    images: List[ParsedImage] = []
    if extract_images:
        output_dir.mkdir(parents=True, exist_ok=True)
        image_idx = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_idx += 1
                image_part = rel.target_part
                ext = Path(image_part.partname).suffix or ".png"
                image_name = f"docx-image-{image_idx:03d}{ext}"
                image_path = output_dir / image_name
                image_path.write_bytes(image_part.blob)
                images.append(
                    ParsedImage(
                        path=str(image_path),
                        page_number=1,  # DOCX is treated as single page
                        image_index=image_idx,
                        bbox=None,  # No spatial info in DOCX
                    )
                )

    page = ParsedPage(
        page_number=1,
        text="\n\n".join(text_parts).strip(),
        blocks=blocks,
        images=images,
        metadata={"source_path": str(docx_path)},
    )

    return ParsedDocument(
        source_path=str(docx_path),
        doc_id=doc_id,
        pages=[page],
        metadata={"parser": "python-docx", "page_count": 1},
    )
